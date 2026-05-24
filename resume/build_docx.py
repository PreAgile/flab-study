#!/usr/bin/env python3
"""
resume_v12.md → resume_v12.docx
- Body font: Pretendard (한글/영문)
- Mono font: D2Coding (코드/모노스페이스)
- 마진 좌우 2cm / 상하 2.2cm
- 이미지: 가로 16cm
- 페이지 번호: 푸터
- 사례 시작 시 page break
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(ROOT, "resume_v12.md")
OUT_PATH = os.path.join(ROOT, "resume_v12.docx")

FONT_BODY = "Pretendard"
FONT_BODY_FALLBACK = "Noto Sans KR"
FONT_MONO = "D2Coding"
FONT_MONO_FALLBACK = "Consolas"

COLOR_BODY = RGBColor(0x1f, 0x29, 0x37)        # slate-800
COLOR_SUBTLE = RGBColor(0x64, 0x74, 0x8b)      # slate-500
COLOR_ACCENT_BLUE = RGBColor(0x1d, 0x4e, 0xd8)
COLOR_ACCENT_GREEN = RGBColor(0x15, 0x80, 0x3d)
COLOR_ACCENT_RED = RGBColor(0xb9, 0x1c, 0x1c)
COLOR_ACCENT_PURPLE = RGBColor(0x70, 0x1a, 0x75)
COLOR_HEADER_PRIMARY = RGBColor(0x0f, 0x17, 0x2a)


def set_run_font(run, size=10.5, bold=False, color=None, mono=False, italic=False):
    run.font.name = FONT_MONO if mono else FONT_BODY
    # East Asia 폰트 별도 지정
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    font = FONT_MONO if mono else FONT_BODY
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:eastAsia'), font)
    rFonts.set(qn('w:cs'), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(p, before=2, after=2, line=1.25):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_horizontal_line(p):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B0BEC5')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def setup_section(doc):
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.gutter = Cm(0)
    # 페이지 번호 (footer)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # PAGE field
    run = fp.add_run()
    set_run_font(run, size=9, color=COLOR_SUBTLE)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = '1'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)


def set_default_style(doc):
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), FONT_BODY)
    rFonts.set(qn('w:hAnsi'), FONT_BODY)
    rFonts.set(qn('w:eastAsia'), FONT_BODY)


def parse_inline(text):
    """Inline 마크다운 파싱 — bold(**), italic(*), code(`), 결과는 (text, attrs) 리스트"""
    tokens = []
    i = 0
    while i < len(text):
        # **bold**
        if text[i:i+2] == "**":
            end = text.find("**", i + 2)
            if end != -1:
                tokens.append((text[i+2:end], {"bold": True}))
                i = end + 2
                continue
        # `code`
        if text[i] == "`":
            end = text.find("`", i + 1)
            if end != -1:
                tokens.append((text[i+1:end], {"mono": True}))
                i = end + 1
                continue
        # plain — accumulate until next marker
        j = i
        while j < len(text) and text[j] != "`" and text[j:j+2] != "**":
            j += 1
        if j > i:
            tokens.append((text[i:j], {}))
        i = j
    return tokens


def add_inline(p, text, base_size=10.5, base_color=None):
    for content, attrs in parse_inline(text):
        if not content:
            continue
        run = p.add_run(content)
        set_run_font(
            run,
            size=base_size,
            bold=attrs.get("bold", False),
            mono=attrs.get("mono", False),
            color=base_color if base_color else COLOR_BODY,
        )


def add_header(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    if level == 1:
        # 도메인 / 메이저 섹션
        set_paragraph_spacing(p, before=14, after=6, line=1.15)
        run = p.add_run(text)
        set_run_font(run, size=15, bold=True, color=COLOR_HEADER_PRIMARY)
        add_horizontal_line(p)
    elif level == 2:
        # 사례 헤딩
        set_paragraph_spacing(p, before=10, after=4, line=1.15)
        run = p.add_run(text)
        set_run_font(run, size=12.5, bold=True, color=COLOR_HEADER_PRIMARY)
    else:
        set_paragraph_spacing(p, before=6, after=2, line=1.15)
        run = p.add_run(text)
        set_run_font(run, size=11.5, bold=True, color=COLOR_HEADER_PRIMARY)


def add_paragraph(doc, text, size=10.5):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=4, line=1.4)
    add_inline(p, text, base_size=size)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.left_indent = Cm(0.6 + level * 0.5)
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.35
    add_inline(p, text, base_size=10.5)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'CFD8DC')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_table_from_md(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=2, after=2, line=1.2)
        add_inline(p, h, base_size=9.5)
        for r in p.runs:
            r.font.bold = True
        shade_cell(cell, "EFF2F6")
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 본문
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[1 + ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=1, after=1, line=1.3)
            add_inline(p, val, base_size=9.5)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 표 후 간격
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=4, line=1.0)


def add_image(doc, png_path, width_cm=16.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=8, after=2, line=1.0)
    run = p.add_run()
    run.add_picture(png_path, width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(cp, before=0, after=8, line=1.0)
        r = cp.add_run(caption)
        set_run_font(r, size=9, color=COLOR_SUBTLE, italic=True)


def add_summary_box(doc, label, value, color):
    """핵심 지표 — 작은 박스 형태 한 줄"""
    pass  # 표로 처리


# ===== Markdown 파서 =====
def parse_md(path):
    with open(path, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")
    return lines


def is_table_separator(line):
    s = line.strip()
    if not s.startswith("|"):
        return False
    return bool(re.match(r"^\|[\s:|-]+\|$", s))


def parse_table_row(line):
    # | a | b | c |
    parts = [p.strip() for p in line.strip().strip("|").split("|")]
    return parts


def build():
    doc = Document()
    set_default_style(doc)
    setup_section(doc)

    lines = parse_md(MD_PATH)
    i = 0
    case_count = 0  # 사례마다 page break
    seen_first_case = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 빈 줄
        if not stripped:
            i += 1
            continue

        # 수평선 — 건너뜀
        if stripped == "---":
            i += 1
            continue

        # 헤더
        if stripped.startswith("# "):
            # 표지 H1 — 이름
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=0, after=2, line=1.1)
            txt = stripped[2:]
            # 「김면수 ｜ Backend Engineer」 split
            run = p.add_run(txt)
            set_run_font(run, size=24, bold=True, color=COLOR_HEADER_PRIMARY)
            i += 1
            continue

        if stripped.startswith("## "):
            heading = stripped[3:]
            # 사례 카드: 「### 사례 N.」 가 아닌 「## 도메인」 같은 메이저 섹션
            # 사례 페이지 분할 정책: 메이저 섹션(도메인 / JVM 재설계 / 경력 / 부록 / 기술 스택 / 교육) 앞 page break
            major_break_keywords = ["도메인 ", "JVM ", "경력 — 아이브릭스", "부록", "기술 스택", "교육"]
            if seen_first_case and any(heading.startswith(k) for k in major_break_keywords):
                add_page_break(doc)
            add_header(doc, heading, level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            heading = stripped[4:]
            # 사례 헤딩 — 첫 번째 사례 이후로는 사례마다 page break
            if heading.startswith("사례 "):
                if seen_first_case:
                    add_page_break(doc)
                else:
                    seen_first_case = True
            add_header(doc, heading, level=2)
            i += 1
            continue

        # 이미지: ![alt](path)
        m = re.match(r"^\!\[([^\]]*)\]\(([^\)]+)\)\s*$", stripped)
        if m:
            alt, path = m.group(1), m.group(2)
            # svg → png 자동 치환
            full = os.path.join(ROOT, path)
            if path.endswith(".svg"):
                png_path = full.replace(".svg", ".png")
            else:
                png_path = full
            if os.path.exists(png_path):
                add_image(doc, png_path, width_cm=16.5, caption=alt)
            else:
                add_paragraph(doc, f"[이미지 누락: {path}]")
            i += 1
            continue

        # 표
        if stripped.startswith("|") and (i + 1) < len(lines) and is_table_separator(lines[i + 1]):
            headers = parse_table_row(stripped)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|") and not is_table_separator(lines[i]):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table_from_md(doc, headers, rows)
            continue

        # 불릿
        if stripped.startswith("- "):
            text = stripped[2:]
            add_bullet(doc, text)
            i += 1
            continue

        # 인용 (>)
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=4, after=4, line=1.35)
            add_inline(p, stripped[2:], base_size=10)
            for r in p.runs:
                r.font.italic = True
                r.font.color.rgb = COLOR_SUBTLE
            i += 1
            continue

        # 일반 문단
        add_paragraph(doc, stripped)
        i += 1

    doc.save(OUT_PATH)
    print(f"✓ saved: {OUT_PATH}  ({os.path.getsize(OUT_PATH)//1024} KB)")


if __name__ == "__main__":
    build()
